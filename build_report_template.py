"""Build templates/report.docx — 5-page A4 landscape Leistungsdiagnostik report.

Round-4 layout (Anna 2026-05-18, supersedes Round 3):
  Page 1: Deckblatt — Logo, Titel, Name, Datum/Ort, kleine Kontaktzeile.
          KEIN "Erstellt von …".
  Page 2: Athletendaten + Testprotokoll (mit Ruhelaktat, Steigung, Dauer letzte
          Stufe) + Rohdaten-Tabelle full-width + Plot full-width (groß, ~25 cm).
  Page 3: Schwellenschnittpunkte (Pivot, post-render, Word-Whitelist) +
          Trainingsbereiche (zonenfarbig, mit Methoden-Spalte). Round 4: keine
          separate Trainingsformen-Mini-Tabelle mehr — Methode ist Spalte.
  Page 4: Interpretation in 4 Blöcken — Zusammenfassung, Schwellen & Zonen,
          Empfehlungen, Energie & Regeneration.
  Page 5: Trainerseite intern — 2×2-Kachelgrid (Pflichtprüfungen, Risiko,
          Schwellenlogik, Testqualität) + Trainernotizen-Tabelle.

Round 4 deltas vs Round 3:
  - Plot deutlich größer (25 cm full-width statt 13 cm side-by-side).
  - Trainingsbereiche bekommt Methoden-Spalte; Mini-Tabelle entfällt.
  - Z1==Z2-Zeile rendert leere Zellen (kein "—") für Intensität/Pace/HF.

The pivot threshold table is built in `src/ld/report.py` after docxtpl renders,
because its column count varies with the number of valid laktat targets.
Zone-row coloring also happens in that post-render pass — docxtpl can't do
conditional per-row cell shading cleanly.

Brand:
  Primary  Dunkelblau   #0B2545
  Accent   Türkis       #16C5A5
  Text     Dunkelgrau   #2C2C2C
  Footer   Mittelgrau   #6B6B6B  (ruhiger Markenanker, nicht Body-Grau)
  Lines    Hellgrau     #D1D5DB
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path("templates/report.docx")
LOGO = Path("assets/logo.jpeg")  # Anna 2026-05-17 (Round 3) — real logo replaces placeholder.
OUT.parent.mkdir(exist_ok=True)

# ── Brand tokens ─────────────────────────────────────────────────────────────
PRIMARY = RGBColor(0x0B, 0x25, 0x45)
ACCENT  = RGBColor(0x16, 0xC5, 0xA5)
TEXT    = RGBColor(0x2C, 0x2C, 0x2C)
FOOTER_GREY = RGBColor(0x6B, 0x6B, 0x6B)
LINE    = "D1D5DB"
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

# Trainer contact (Anna 2026-05-13)
CONTACT_EMAIL = "anna-maria@woerndle.at"
CONTACT_PHONE = "+43 677 62150496"


# ── Document setup: A4 landscape, denser margins for the 5-page budget ──────
doc = Document()

styles = doc.styles
for sname in ("Normal",):
    style = styles[sname]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = TEXT

# Set spacing-after on Normal to 4pt — base "tighter than Word default".
pPr = styles["Normal"].element.get_or_add_pPr()
spacing = OxmlElement("w:spacing")
spacing.set(qn("w:after"), "80")  # 80 twips ~= 4pt
spacing.set(qn("w:line"), "260")  # ~1.05 line height
spacing.set(qn("w:lineRule"), "auto")
pPr.append(spacing)

section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width  = Cm(29.7)
section.page_height = Cm(21.0)
section.top_margin    = Cm(1.3)
section.bottom_margin = Cm(1.4)
section.left_margin   = Cm(1.6)
section.right_margin  = Cm(1.6)
section.different_first_page_header_footer = True


# ── Helpers ──────────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_no_hash: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_no_hash)
    tc_pr.append(shd)


def _set_cell_borders(cell, color_hex: str = LINE, size_pt: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size_pt))
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tc_pr.append(borders)


def _set_cell_no_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tc_pr.append(borders)


def heading(text: str, level: int = 1, color: RGBColor = PRIMARY,
            size_pt: int | None = None) -> None:
    p = doc.add_heading(text, level=level)
    if p.runs:
        run = p.runs[0]
        run.font.color.rgb = color
        run.font.size = Pt(size_pt or {1: 15, 2: 12, 3: 10}.get(level, 10))
        run.font.bold = True


def body(text: str, *, bold: bool = False, color: RGBColor = TEXT,
         size_pt: int = 10, italic: bool = False) -> None:
    p = doc.add_paragraph(text)
    if p.runs:
        run = p.runs[0]
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = italic


def page_break() -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_table_with_loop(*, headers: list[str], row_template: list[str], loop_var: str,
                        items_var: str, header_widths_cm: list[float] | None = None,
                        header_font_size: int = 9, body_font_size: int = 9) -> None:
    """Single header + docxtpl row-loop body (1 visible row, repeated)."""
    t = doc.add_table(rows=4, cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.autofit = False
    if header_widths_cm:
        for i, w in enumerate(header_widths_cm):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        cell.text = ""
        _set_cell_bg(cell, "0B2545")
        _set_cell_borders(cell, "0B2545")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.font.bold = True
        run.font.size = Pt(header_font_size)
        run.font.color.rgb = WHITE

    tag_open = t.cell(1, 0)
    tag_open.text = ""
    tag_open.paragraphs[0].add_run("{%tr for " + loop_var + " in " + items_var + " %}")

    for i, tpl in enumerate(row_template):
        cell = t.cell(2, i)
        cell.text = ""
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(tpl)
        run.font.size = Pt(body_font_size)
        run.font.color.rgb = TEXT

    tag_close = t.cell(3, 0)
    tag_close.text = ""
    tag_close.paragraphs[0].add_run("{%tr endfor %}")


def add_kv_table(rows: list[tuple[str, str]], *, key_width_cm: float = 4.5,
                 val_width_cm: float = 7.5, font_size: int = 9) -> None:
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    for i, (k, v) in enumerate(rows):
        kc, vc = t.cell(i, 0), t.cell(i, 1)
        kc.width = Cm(key_width_cm); vc.width = Cm(val_width_cm)
        _set_cell_borders(kc); _set_cell_borders(vc)
        kc.text = ""; vc.text = ""
        kr = kc.paragraphs[0].add_run(k)
        kr.font.bold = True; kr.font.size = Pt(font_size); kr.font.color.rgb = PRIMARY
        vr = vc.paragraphs[0].add_run(v)
        vr.font.size = Pt(font_size); vr.font.color.rgb = TEXT


def add_two_column_kv(left_rows: list[tuple[str, str]],
                      right_rows: list[tuple[str, str]],
                      *, left_title: str, right_title: str) -> None:
    """Render two key-value tables side-by-side inside a 2x1 outer table.
    Used on Page 2 to fit Athletendaten + Testprotokoll on the same vertical band.
    """
    outer = doc.add_table(rows=2, cols=2)
    outer.autofit = False
    for col, w in enumerate((13.5, 13.5)):
        for r in outer.rows:
            r.cells[col].width = Cm(w)
    for c in [outer.cell(r, c) for r in range(2) for c in range(2)]:
        _set_cell_no_borders(c)

    # Titles in row 0.
    for col, title in enumerate((left_title, right_title)):
        cell = outer.cell(0, col)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(title)
        run.font.bold = True; run.font.size = Pt(12); run.font.color.rgb = PRIMARY

    # Inner KV-tables in row 1, one per side. Round 4 (Anna 2026-05-18): 8 pt
    # statt 8.5 pt + tight spacing — Page 2 muss Daten-Block + Rohdaten +
    # voller Plot tragen, jeder halbe Millimeter Zeilenhöhe zählt.
    for col, rows in enumerate((left_rows, right_rows)):
        host_cell = outer.cell(1, col)
        # Replace placeholder paragraph with a sub-table.
        para = host_cell.paragraphs[0]
        para.text = ""
        inner = host_cell.add_table(rows=len(rows), cols=2)
        inner.autofit = False
        for i, (k, v) in enumerate(rows):
            kc, vc = inner.cell(i, 0), inner.cell(i, 1)
            kc.width = Cm(4.2); vc.width = Cm(8.5)
            _set_cell_borders(kc); _set_cell_borders(vc)
            kc.text = ""; vc.text = ""
            # Tight paragraph spacing — no space before/after — so 14 KV rows
            # don't eat the height budget. python-docx default spacing is
            # ~80 twips after each paragraph.
            for cell in (kc, vc):
                p_pPr = cell.paragraphs[0]._element.get_or_add_pPr()
                p_spacing = OxmlElement("w:spacing")
                p_spacing.set(qn("w:before"), "0")
                p_spacing.set(qn("w:after"), "0")
                p_spacing.set(qn("w:line"), "240")  # single line height
                p_spacing.set(qn("w:lineRule"), "auto")
                p_pPr.append(p_spacing)
            kr = kc.paragraphs[0].add_run(k)
            kr.font.bold = True; kr.font.size = Pt(8); kr.font.color.rgb = PRIMARY
            vr = vc.paragraphs[0].add_run(v)
            vr.font.size = Pt(8); vr.font.color.rgb = TEXT
        # Remove the empty placeholder paragraph python-docx added inside the host cell.
        if host_cell.paragraphs and not host_cell.paragraphs[0].text:
            p_elem = host_cell.paragraphs[0]._element
            p_elem.getparent().remove(p_elem)


def add_2x2_grid(quadrants: list[tuple[str, str]]) -> None:
    """A simple 2x2 grid for Page 5's internal coaching view.
    Each quadrant: (title, body)."""
    assert len(quadrants) == 4
    t = doc.add_table(rows=2, cols=2)
    t.autofit = False
    for col in range(2):
        for r in t.rows:
            r.cells[col].width = Cm(13.5)
    cells = [t.cell(r, c) for r in range(2) for c in range(2)]
    for cell, (title, body_text) in zip(cells, quadrants):
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Title
        p = cell.paragraphs[0]
        p.text = ""
        run = p.add_run(title)
        run.font.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = PRIMARY
        # Body — may include Jinja placeholders, so emit as separate paragraph
        for line in body_text.split("\n"):
            bp = cell.add_paragraph()
            br = bp.add_run(line)
            br.font.size = Pt(9); br.font.color.rgb = TEXT


# ── PAGE 1: Deckblatt ────────────────────────────────────────────────────────
# Round 2: airy cover, Anna-Maria-line removed, Sport-AnnaLytics-Subtext added.
for _ in range(2):
    doc.add_paragraph()

if LOGO.exists():
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.add_run().add_picture(str(LOGO), width=Cm(6.0))

for _ in range(2):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Leistungsdiagnostik {{ athlete.sportart_label }}")
title_run.font.size = Pt(32); title_run.font.bold = True; title_run.font.color.rgb = PRIMARY

brand_p = doc.add_paragraph()
brand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
brand_run = brand_p.add_run("Sport AnnaLytics")
brand_run.font.size = Pt(13); brand_run.font.color.rgb = ACCENT
brand_run.font.italic = False

for _ in range(2):
    doc.add_paragraph()

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name_p.add_run("{{ athlete.vorname }} {{ athlete.name }}")
name_run.font.size = Pt(22); name_run.font.color.rgb = TEXT

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta_p.add_run("{{ testprotokoll.testdatum }}  ·  {{ testprotokoll.durchfuehrungsort }}")
meta_run.font.size = Pt(12); meta_run.font.color.rgb = TEXT

# Round 3 P0-3 (Anna 2026-05-17): Kontaktzeile unten am Deckblatt — klein,
# zurückhaltend. Deckblatt bleibt ruhig/premium, aber Kontakt darf erscheinen
# (Deckblatt wird oft separat geteilt/gedruckt). "Erstellt von..." bleibt entfernt.
for _ in range(3):
    doc.add_paragraph()

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact_p.add_run(f"{CONTACT_EMAIL}  |  {CONTACT_PHONE}")
contact_run.font.size = Pt(9); contact_run.font.color.rgb = FOOTER_GREY


# ── PAGE 2: Athletendaten + Testprotokoll + Rohdaten + Plot ─────────────────
page_break()
heading("Daten & Testablauf", level=1, size_pt=14)

add_two_column_kv(
    left_title="Athletendaten",
    left_rows=[
        ("Name",            "{{ athlete.vorname }} {{ athlete.name }}"),
        ("Email",           "{{ athlete.email }}"),
        ("Geburtsjahr",     "{{ athlete.geburtsjahr }}"),
        ("Alter",           "{{ athlete.alter }}"),
        ("Geschlecht",      "{{ athlete.geschlecht }}"),
        ("Gewicht",         "{{ athlete.gewicht_kg }} kg"),
        ("Größe",           "{{ athlete.groesse_m }} m"),
        ("Trainingsziel",   "{{ athlete.trainingsziel }}"),
        ("Wettkampfziel",   "{{ athlete.wettkampfziel }}"),
        ("Umfang/Woche",    "{{ athlete.trainingsumfang_wo }}"),
        ("Leistungsniveau", "{{ athlete.leistungsniveau }}"),
    ],
    right_title="Testprotokoll",
    right_rows=[
        ("Sportart",        "{{ testprotokoll.sportart_label }}"),
        ("Datum",           "{{ testprotokoll.testdatum }}"),
        ("Uhrzeit",         "{{ testprotokoll.uhrzeit }}"),
        ("Ort",             "{{ testprotokoll.durchfuehrungsort }}"),
        ("Testleiter",      "{{ testprotokoll.testleiter }}"),
        ("Gerät",           "{{ testprotokoll.geraet }}"),
        ("Anfangsbelastung","{{ testprotokoll.anfangsbelastung_display }}"),
        ("Stufeninkrement", "{{ testprotokoll.stufeninkrement_display }}"),
        ("Stufendauer",     "{{ testprotokoll.stufendauer_min }} min"),
        # Round 3 P0-2 (Anna 2026-05-17): Ruhelaktat, Steigung, Dauer letzte Stufe.
        ("Ruhelaktat",      "{{ testprotokoll.ruhelaktat_display }}"),
        ("Steigung",        "{{ testprotokoll.steigung_display }}"),
        ("Dauer letzte Stufe", "{{ testprotokoll.dauer_letzte_stufe_display }}"),
        ("Ausbelastung",    "{{ ausbelastung_de }}"),
        ("{{ v_max_label }}", "{{ v_max_display }}"),
    ],
)

# Round 4 (Anna 2026-05-18): Plot deutlich größer. Statt side-by-side mit der
# Rohdaten-Tabelle steht die Tabelle jetzt full-width schmal über dem
# full-width Plot. So bekommt der Plot ~25 cm Breite statt 13 cm und wirkt
# als primäres visuelles Element der Seite — Anna explizit: "Grafikgröße
# hat Vorrang vor der starren Annahme, dass die Grafik zwingend auf Seite 2
# bleiben muss". Wenn das 5-Seiten-Budget bricht, ist der Fallback Variante 2
# (Plot rückt auf Seite 3).
# Eine kombinierte Überschrift statt zwei separaten — spart vertikale
# Höhe, die wir für den größeren Plot brauchen. e2e-Test erwartet sowohl
# "Rohdaten" als auch "Diagramm" als Anker auf Seite 2.
heading("Rohdaten der Teststufen & Diagramm", level=3, size_pt=11)
raw = doc.add_table(rows=4, cols=5)
raw.autofit = False
raw.alignment = WD_ALIGN_PARAGRAPH.CENTER
raw_widths = [2.4, 5.3, 5.3, 6.3, 3.7]   # ~23 cm gesamt, kompakt
for i, w in enumerate(raw_widths):
    for r in raw.rows:
        r.cells[i].width = Cm(w)
raw_headers = ["Stufe", "{{ x_axis_label }}", "HF (bpm)", "Laktat (mmol/l)", "RPE (0-10)"]
for i, h in enumerate(raw_headers):
    c = raw.cell(0, i)
    c.text = ""
    _set_cell_bg(c, "0B2545"); _set_cell_borders(c, "0B2545")
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h); run.font.bold = True; run.font.size = Pt(8.5); run.font.color.rgb = WHITE
raw.cell(1, 0).text = ""
raw.cell(1, 0).paragraphs[0].add_run("{%tr for s in steps_display %}")
row_tpl = ["{{ s.stufe }}", "{{ s.intensitaet }}", "{{ s.herzfrequenz_bpm }}",
           "{{ s.laktat_mmol }}", "{{ s.rpe }}"]
for i, tpl in enumerate(row_tpl):
    c = raw.cell(2, i)
    c.text = ""
    _set_cell_borders(c)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(tpl); run.font.size = Pt(8.5); run.font.color.rgb = TEXT
raw.cell(3, 0).text = ""
raw.cell(3, 0).paragraphs[0].add_run("{%tr endfor %}")

# Plot full-width, zentriert, direkt unter der kombinierten
# Rohdaten/Diagramm-Überschrift.
plot_p = doc.add_paragraph()
plot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
plot_p.add_run("{{ diagram }}")


# ── PAGE 3: Analyse (Pivot + Zonen) ─────────────────────────────────────────
page_break()
heading("Analyse der Leistungsdiagnostik", level=1, size_pt=14)

heading("Schwellenschnittpunkte", level=2)
# This paragraph is replaced after docxtpl render by a dynamically-built pivot
# table. See _insert_intersection_pivot() in src/ld/report.py.
marker_p = doc.add_paragraph("<<INTERSECTION_PIVOT>>")
mr = marker_p.runs[0]; mr.font.color.rgb = TEXT; mr.font.size = Pt(9)

doc.add_paragraph()
heading("Trainingsbereiche", level=2)
# Rows are colored per zone in the post-render pass (_color_zone_rows).
# Round 4 (Anna 2026-05-18): Methoden-Spalte ist Teil der Tabelle (Variante A
# aus Round 3, früher verworfen). Separate Mini-Tabelle entfällt.
# 7 Spalten — Methode bekommt großzügigste Breite weil längster Text.
add_table_with_loop(
    headers=["Zone", "Ziel", "{{ x_axis_label }}", "Pace (min/km)",
             "Herzfrequenz (bpm)", "RPE (0-10)", "Methode / Trainingsform"],
    row_template=[
        "{{ z.name }}",
        "{{ z.ziel }}",
        "{{ z.intensitaet_range }}",
        "{{ z.pace_range }}",
        "{{ z.herzfrequenz_range }}",
        "{{ z.rpe }}",
        "{{ z.methode }}",
    ],
    loop_var="z",
    items_var="zones",
    header_widths_cm=[1.3, 3.4, 3.4, 3.4, 3.6, 1.8, 5.8],
    header_font_size=9,
    body_font_size=9,
)


# ── PAGE 4: Interpretation (4 Blöcke) ───────────────────────────────────────
page_break()
heading("Interpretation", level=1, size_pt=14)

heading("Zusammenfassung", level=2)
body("{{ interp_zusammenfassung }}")

heading("Schwellen & Zonen", level=2)
body("{{ interp_schwellen }}")

# Round 3 P0-9 (Anna 2026-05-17): "Nächste 3-4 Wochen" → "Empfehlungen".
# Beispielwoche wird auf Prompt-Ebene entfernt (siehe .codex/prompts/ld-report.md).
heading("Empfehlungen", level=2)
body("{{ interp_coaching_ausblick }}")

heading("Energie & Regeneration", level=2)
body("{{ interp_ernaehrung }}")


# ── PAGE 5: Trainerseite intern (2×2 + Trainernotizen-Table) ────────────────
page_break()
intern_p = doc.add_paragraph()
intern_run = intern_p.add_run("Trainerseite — Intern (nicht für Athlet:in)")
intern_run.font.size = Pt(10); intern_run.font.italic = True; intern_run.font.color.rgb = ACCENT
heading("Fachliche Notizen", level=1, size_pt=14)

# 2x2 quadrants — bodies are plain template variables (conditional logic and
# loops are pre-formatted in src/ld/report.py to avoid docxtpl paragraph quirks).
add_2x2_grid([
    ("Pflichtprüfungen", "{{ pflichtpruefungen_text }}"),
    (
        "Risikoanalyse",
        "{{ interp_risiko if interp_risiko else 'Keine besonderen Risiken aus den Daten ableitbar. Siehe Pflichtprüfungen für Detailmuster.' }}",
    ),
    (
        "Schwellenlogik",
        "Zonengrenzen sind Orientierungspunkte. Z2-Ober: v bei 2.0 mmol/l. "
        "Z3-Ober: v bei 3.0 mmol/l. Z4-Ober: v bei 4.0 mmol/l. "
        "Z5-Ober: Maximalgeschwindigkeit (aliquot-korrigiert). Z6 darüber. "
        "Diese Werte werden mit Kurvenform, HF-Verlauf und RPE-Muster kombiniert "
        "(keine rein fixen mmol-Schwellen).",
    ),
    ("Testqualität", "{{ testqualitaet_text }}"),
])

doc.add_paragraph()
heading("Trainernotizen", level=2)
add_kv_table([
    ("Verletzungen",          "{{ coaching.verletzungen }}"),
    ("Aktuelle Probleme",     "{{ coaching.aktuelle_probleme }}"),
    ("Stärken",               "{{ coaching.staerken }}"),
    ("Schwächen",             "{{ coaching.schwaechen }}"),
    ("Geplante Wettkämpfe",   "{{ coaching.geplante_wettkaempfe }}"),
    ("Notizen",               "{{ coaching.trainernotizen }}"),
], font_size=9)


# ── Footer (pages 2-5, suppressed on page 1) ────────────────────────────────
def _build_footer(section, *, primary_color: RGBColor = PRIMARY) -> None:
    """Footer mit durchgehender Trennlinie (oben am Footer-Bereich, nicht
    pro Zelle — sonst sieht es segmentiert aus). Mittelgrau für ruhigen
    Markenanker (Round 2 P1-2)."""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    # Trennlinie als Border am ersten Paragraphen — durchgehende Linie über
    # die volle Breite, weil ein Paragraph keine Cell-Grenzen hat.
    rule_p = footer.add_paragraph()
    pPr = rule_p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), LINE)
    pBdr.append(top)
    pPr.append(pBdr)

    table = footer.add_table(rows=1, cols=3, width=Cm(26.5))
    table.autofit = False
    cells = table.rows[0].cells
    cells[0].width = Cm(12.0); cells[1].width = Cm(8.0); cells[2].width = Cm(6.5)
    for c in cells:
        _set_cell_no_borders(c)

    # Left: Kontakt — Round 3 P0-10 (Anna 2026-05-17): zwei Zeilen untereinander
    # statt "email · phone". Wirkt ruhiger und professioneller. Erster
    # Paragraph nimmt die obere Trennlinie via pBdr; Email steht im zweiten
    # Paragraph der Zelle, Telefon im dritten.
    cl_email = cells[0].paragraphs[0]
    cl_email.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_email = cl_email.add_run(CONTACT_EMAIL)
    r_email.font.size = Pt(8); r_email.font.color.rgb = FOOTER_GREY

    cl_phone = cells[0].add_paragraph()
    cl_phone.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Tighten spacing between the two lines.
    cl_phone_pPr = cl_phone._element.get_or_add_pPr()
    cl_phone_spacing = OxmlElement("w:spacing")
    cl_phone_spacing.set(qn("w:before"), "0")
    cl_phone_spacing.set(qn("w:after"), "0")
    cl_phone_spacing.set(qn("w:line"), "240")
    cl_phone_spacing.set(qn("w:lineRule"), "auto")
    cl_phone_pPr.append(cl_phone_spacing)
    r_phone = cl_phone.add_run(CONTACT_PHONE)
    r_phone.font.size = Pt(8); r_phone.font.color.rgb = FOOTER_GREY

    # Center: Seitenzahl.
    cm = cells[1].paragraphs[0]
    cm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cm.add_run("Seite ")
    run.font.size = Pt(8); run.font.color.rgb = FOOTER_GREY

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    page_run = cm.add_run()
    page_run.font.size = Pt(8); page_run.font.color.rgb = FOOTER_GREY
    page_run._r.append(fld_char_begin); page_run._r.append(instr); page_run._r.append(fld_char_end)

    # Right: Mini-Logo.
    cr = cells[2].paragraphs[0]
    cr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO.exists():
        cr.add_run().add_picture(str(LOGO), width=Cm(2.0))


def _hide_first_page_footer(section) -> None:
    fp = section.first_page_footer
    for p in list(fp.paragraphs):
        p._element.getparent().remove(p._element)
    fp.add_paragraph()


_build_footer(section)
_hide_first_page_footer(section)


# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Report-Template gespeichert: {OUT}")
