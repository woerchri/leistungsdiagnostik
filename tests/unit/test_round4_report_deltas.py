"""Round 4 (Anna 2026-05-18) — render-level smoke tests for the deltas vs Round 3.

Round-4 changes:
  - Z1==Z2 collapse renders TRULY EMPTY cells (not "—") for Intensität/Pace/HF
    in the Trainingsbereiche table. (P0-2)
  - Trainingsbereiche table has a "Methode / Trainingsform" column. (P0-4)
  - Separate "Beschreibung / Methode Trainingsformen" mini-table is GONE.
  - Plot is significantly larger on Page 2 (full-width below Rohdaten,
    not side-by-side). (P0-3) — asserted indirectly via Page-2 layout test.
  - Codex prompt requires real Umlauts (ä/ö/ü/ß); no ae/oe/ue/ss substitutes
    in customer-visible interpretation prose. (P0-6) — verified at the prompt
    level (tests/unit/test_codex_prompt.py-like check).
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

import docx2txt
from docx import Document

from ld import io_input, plots, protocols, report, zones
from ld.types import (
    Athlete,
    Coaching,
    TestRun,
    TestStep,
    Testprotokoll,
)


FIXTURE = Path(__file__).parent.parent / "protocols" / "fixtures" / "lauf" / "rainier.xlsx"


def _render_rainier(tmp_path):
    run = io_input.parse_input(FIXTURE)
    result = protocols.analyze(run)
    plot_path = plots.render_main_diagram(result, tmp_path / "plots")
    out = tmp_path / "out.docx"
    report.render(result, plot_path, out, interpretation=None)
    return result, out


# ── P0-4: Trainingsbereiche has Methode column; separate mini-table is gone ──

def test_trainingsbereiche_has_methode_column(tmp_path):
    """Round 4 P0-4: Methode/Trainingsform ist eine SPALTE in
    Trainingsbereiche, keine separate Mini-Tabelle mehr (Variante A statt B
    aus Round 3)."""
    _, docx_path = _render_rainier(tmp_path)
    doc = Document(str(docx_path))

    zones_table = None
    for table in doc.tables:
        if not table.rows:
            continue
        header_texts = [c.text.strip() for c in table.rows[0].cells]
        if "Zone" in header_texts and "Ziel" in header_texts:
            zones_table = table
            break
    assert zones_table is not None, "Trainingsbereiche-Tabelle nicht gefunden"

    header_texts = [c.text.strip() for c in zones_table.rows[0].cells]
    assert "Methode / Trainingsform" in header_texts, (
        f"Methode-Spalte fehlt — Header sind: {header_texts}"
    )
    # 7-Spalten-Layout (Round 3 hatte 6): Zone | Ziel | Intensität | Pace | HF | RPE | Methode
    assert len(header_texts) == 7, (
        f"Trainingsbereiche-Tabelle sollte 7 Spalten haben, hat {len(header_texts)}: {header_texts}"
    )

    # Methoden-Inhalt in den Zone-Zeilen — alle 6 Round-4-Phrasings müssen vorkommen.
    methode_col_idx = header_texts.index("Methode / Trainingsform")
    method_cells_by_zone = {
        row.cells[0].text.strip(): row.cells[methode_col_idx].text.strip()
        for row in zones_table.rows[1:]
        if row.cells and row.cells[0].text.strip().startswith("Z")
    }
    for zone_name, expected in zones.ZONE_METHODE.items():
        assert method_cells_by_zone.get(zone_name) == expected, (
            f"{zone_name} Methode-Zelle: erwartet {expected!r}, "
            f"gerendert {method_cells_by_zone.get(zone_name)!r}"
        )


def test_separate_trainingsformen_mini_table_removed(tmp_path):
    """Round 4 P0-4: alte Round-3-Mini-Tabelle mit Überschrift
    'Beschreibung / Methode Trainingsformen' wurde entfernt — diese Phrase
    darf nirgendwo mehr im Wordbericht stehen."""
    _, docx_path = _render_rainier(tmp_path)
    text = docx2txt.process(str(docx_path))
    assert "Beschreibung / Methode Trainingsformen" not in text, (
        "Alte Round-3-Mini-Tabellen-Überschrift ist noch vorhanden — sollte "
        "durch die Methode-Spalte in Trainingsbereiche ersetzt sein."
    )


# ── P0-2: Z1==Z2 collapse renders empty cells, not "—" ──────────────────────

def _build_z1_collapsed_testrun() -> TestRun:
    """Konstruiert einen synthetischen Lauf-Testlauf, bei dem v(lk=2) und
    v(lk=3) genau zusammenfallen — dann markiert suggest_zones() Z1 als
    is_collapsed_with_z2=True und die Render-Schicht muss leere Zellen
    statt '—' liefern."""
    # Künstliche Stufendaten: Laktat steigt langsam genug, dass die kubische
    # Approximation lk=2 und lk=3 nahe beieinander platziert, plus ein
    # steiler Sprung am Ende, damit der lk=4-Schnittpunkt klar später kommt.
    steps = (
        TestStep(stufe=1, intensitaet=7.0,  herzfrequenz_bpm=130, laktat_mmol=1.0, rpe=2),
        TestStep(stufe=2, intensitaet=8.0,  herzfrequenz_bpm=140, laktat_mmol=1.5, rpe=3),
        TestStep(stufe=3, intensitaet=9.0,  herzfrequenz_bpm=150, laktat_mmol=2.5, rpe=5),
        TestStep(stufe=4, intensitaet=10.0, herzfrequenz_bpm=160, laktat_mmol=3.5, rpe=7),
        TestStep(stufe=5, intensitaet=11.0, herzfrequenz_bpm=170, laktat_mmol=5.5, rpe=8),
        TestStep(stufe=6, intensitaet=12.0, herzfrequenz_bpm=178, laktat_mmol=8.0, rpe=9),
    )
    athlete = Athlete(
        sportart="lauf", vorname="Test", name="Z1Collapse",
        geburtsjahr=1990, geschlecht="m", gewicht_kg=70.0, groesse_m=1.80,
        trainingsziel="Test", wettkampfziel="Test",
        trainingsumfang_wo="5h", leistungsniveau="Hobby",
        email=None,
    )
    proto = Testprotokoll(
        testdatum=date(2026, 5, 18), uhrzeit="10:00",
        durchfuehrungsort="Lab", testleiter="Anna",
        geraet="Laufband", anfangsbelastung=7.0, stufeninkrement=1.0,
        stufendauer_min=4.0, stufenlaenge_m=None, besonderheiten="",
        letzte_stufe_vollstaendig=True, dauer_letzte_stufe_min=None,
        ausbelastung=True,
    )
    coaching = Coaching(
        verletzungen="", aktuelle_probleme="", staerken="",
        schwaechen="", geplante_wettkaempfe="", trainernotizen="",
    )
    return TestRun(athlete=athlete, testprotokoll=proto, steps=steps, coaching=coaching)


def test_z1_collapsed_renders_empty_cells_not_dash(tmp_path):
    """Round 4 P0-2: wenn Z1 mit Z2 zusammenfällt, sind Intensität/Pace/HF
    in der Z1-Zeile WIRKLICH LEER (kein '—'). '—' wirkt auf Athlet:innen wie
    'Wert fehlt', während die leere Zelle 'kein eigener Bereich, siehe Z2'
    bedeutet.

    Dieser Test prüft die Render-Schicht direkt mit einem konstruierten
    Zonen-Set, weil die Rainier-Fixture keinen Z1-Collapse auslöst."""
    from io import BytesIO

    run = _build_z1_collapsed_testrun()
    result = protocols.analyze(run)

    # Falls das synthetische Setup keinen Collapse triggert: sanity-check
    # überspringen — der eigentliche Render-Layer-Test ist unten via Monkey.
    z1 = next(z for z in result.zones_final if z.name == "Z1")
    if not z1.is_collapsed_with_z2:
        # Forciere Collapse über `replace` — testet rein die Render-Schicht.
        from dataclasses import replace
        new_zones = tuple(
            replace(z, is_collapsed_with_z2=True,
                    intensitaet_min=None, intensitaet_max=None,
                    herzfrequenz_min=None, herzfrequenz_max=None,
                    pace_min_min_per_km=None, pace_max_min_per_km=None,
                    is_open_lower=False)
            if z.name == "Z1" else z
            for z in result.zones_final
        )
        result = replace(result, zones_final=new_zones)

    plot_path = plots.render_main_diagram(result, tmp_path / "plots")
    out = tmp_path / "out.docx"
    report.render(result, plot_path, out, interpretation=None)

    doc = Document(str(out))
    zones_table = next(
        t for t in doc.tables
        if t.rows and "Zone" in [c.text.strip() for c in t.rows[0].cells]
           and "Ziel" in [c.text.strip() for c in t.rows[0].cells]
    )
    header = [c.text.strip() for c in zones_table.rows[0].cells]
    z1_row = next(r for r in zones_table.rows[1:]
                  if r.cells and r.cells[0].text.strip() == "Z1")
    cells = [c.text.strip() for c in z1_row.cells]

    # Layout: Zone | Ziel | Intensität | Pace | HF | RPE | Methode
    intens_idx = header.index("Geschwindigkeit (km/h)")
    pace_idx = header.index("Pace (min/km)")
    hf_idx = header.index("Herzfrequenz (bpm)")
    methode_idx = header.index("Methode / Trainingsform")

    assert cells[intens_idx] == "", (
        f"Z1 Intensität sollte LEER sein (Round 4), war {cells[intens_idx]!r}"
    )
    assert cells[pace_idx] == "", (
        f"Z1 Pace sollte LEER sein (Round 4), war {cells[pace_idx]!r}"
    )
    assert cells[hf_idx] == "", (
        f"Z1 HF sollte LEER sein (Round 4), war {cells[hf_idx]!r}"
    )
    # Methode + Ziel + RPE bleiben aber sichtbar — sonst wäre Z1 komplett unsichtbar.
    assert cells[methode_idx] == zones.ZONE_METHODE["Z1"], (
        f"Z1 Methode sollte trotz Collapse sichtbar bleiben, war {cells[methode_idx]!r}"
    )
    assert cells[1] == "Regeneration" or "Regeneration" in cells[1], (
        f"Z1 Ziel sollte 'Regeneration' enthalten, war {cells[1]!r}"
    )


# ── P0-6: Codex prompt mandates real Umlauts ────────────────────────────────

def test_codex_prompt_contains_umlaut_rule():
    """Round 4 P0-6: der Codex-Prompt MUSS eine harte Regel enthalten,
    dass echte Umlaute (ä/ö/ü/ß) verwendet werden und niemals die
    ASCII-Ersatzschreibweise ae/oe/ue/ss in kundensichtbarem Text."""
    prompt_path = Path(__file__).parent.parent.parent / ".codex" / "prompts" / "ld-report.md"
    text = prompt_path.read_text(encoding="utf-8")
    assert "Echte Umlaute" in text or "echte Umlaute" in text, (
        "Codex-Prompt enthält keine explizite Umlaut-Regel — Round 4 P0-6 fehlt."
    )
    # Mindestens ein Negativ-Beispiel ("nicht `fuer`") muss im Prompt stehen,
    # damit klar ist, was Anna ausschließen will.
    assert "fuer" in text and "für" in text, (
        "Codex-Prompt sollte konkrete Negativ-/Positivbeispiele für die "
        "Umlaut-Regel enthalten."
    )
