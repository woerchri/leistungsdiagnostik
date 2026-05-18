from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docxtpl import DocxTemplate, InlineImage

from ld.types import AnalysisResult
from ld.zones import ZONE_METHODE


_TEMPLATE_PATH = Path(__file__).parent.parent.parent / "templates" / "report.docx"

# Round 3 P0-5 (Anna 2026-05-17): kuratierte Whitelist für den Wordbericht.
# Die Pipeline berechnet weiterhin alle Targets (1.0..8.0) in JSON/Codex —
# der Wordbericht zeigt aber NUR diese Werte. 2.5 wird komplett weggelassen
# (laut Anna immer); 8.0 wird zusätzlich noch durch eine "reached"-Prüfung
# gefiltert (Wert nur, wenn er innerhalb des gemessenen Bereichs liegt).
_WORD_REPORT_LAKTAT_WHITELIST: tuple[float, ...] = (1.0, 1.5, 2.0, 3.0, 4.0, 6.0, 8.0)
_WORD_REPORT_LAKTAT_REACHED_ONLY: tuple[float, ...] = (8.0,)
# A4 landscape gives ~25 cm usable width. Anna 2026-05-13 Round 2 P0-2:
# reduce plot width so Page 2 fits athlete + protocol + raw-data table + plot.
_PLOT_WIDTH_CM = 13.0  # Fits in right-side cell of the Page-2 Rohdaten+Plot row.

# Brand tokens — mirrored from build_report_template.py.
_PRIMARY = RGBColor(0x0B, 0x25, 0x45)
_ACCENT = RGBColor(0x16, 0xC5, 0xA5)
_TEXT = RGBColor(0x2C, 0x2C, 0x2C)
_LINE = "D1D5DB"
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Zone-row background fills (Round 2 P1-3 — desaturated/transparent-like tones,
# matching the plot zone bands but at print-safe lightness).
_ZONE_BG: dict[str, str] = {
    "Z1": "E3F0FA",  # Hellblau (light)
    "Z2": "E4F4DC",  # Grün (light)
    "Z3": "FFF6CC",  # Gelb (light)
    "Z4": "F7E9B0",  # Dunkelgelb (light)
    "Z5": "FFE0CC",  # Orange (light)
    "Z6": "FBDBDC",  # Rot (light)
}

# Markers — paragraphs whose .text starts with one of these are placeholders
# that the post-render pass replaces with a dynamically-built table.
_PIVOT_MARKER = "<<INTERSECTION_PIVOT>>"


def _x_axis_label(sport: str) -> str:
    if sport in {"lauf", "triathlon-lauf"}:
        return "Geschwindigkeit (km/h)"
    if sport in {"rad", "triathlon-rad"}:
        return "Leistung (W)"
    return "Stufe"


# Per-field char caps for Page-5 Trainernotizen-Tabelle. Anna 2026-05-17:
# "Kürzen, sodass die Notizen auf einer Seite Platz haben". Truncation here
# protects the hard 5-page cap without losing the diagnostic information —
# the FULL data is still in the JSON for archive purposes.
_COACHING_CHAR_CAPS: dict[str, int] = {
    "verletzungen": 120,
    "aktuelle_probleme": 120,
    "staerken": 120,
    "schwaechen": 120,
    "geplante_wettkaempfe": 120,
    "trainernotizen": 260,  # Free-text — slightly larger budget.
}


def _truncate(text: str | None, max_chars: int, *, field: str) -> str:
    """Trim to `max_chars` minus the ellipsis. Prints a warning to stdout so
    the operator notices that some text was suppressed for the page budget."""
    if not text:
        return "—"
    if len(text) <= max_chars:
        return text
    print(
        f"Hinweis: Trainernotiz-Feld '{field}' ({len(text)} Zeichen) wurde auf "
        f"{max_chars} gekürzt — vollständiger Text bleibt in der JSON-Datei."
    )
    return text[: max_chars - 1].rstrip() + "…"


def _format_pflichtpruefungen(result: AnalysisResult) -> str:
    failed = [p for p in result.pflichtpruefungen if not p.ok]
    total = len(result.pflichtpruefungen)
    passed = total - len(failed)
    if not failed:
        return f"Alle Pflichtprüfungen: OK ({passed} / {total})."
    lines = ["Hinweise vor der Interpretation:"]
    for p in failed:
        lines.append(f"• {p.message_de}")
    return "\n".join(lines)


def _format_testqualitaet(result: AnalysisResult) -> str:
    total = len(result.pflichtpruefungen)
    passed = sum(1 for p in result.pflichtpruefungen if p.ok)
    base = f"Pflichtprüfungen bestanden: {passed} von {total}."
    blocking = [p for p in result.pflichtpruefungen
                if not p.ok and p.name in {"letzte_stufe", "ausbelastung"}]
    if blocking:
        base += (
            " Hinweis: " + ", ".join(p.name for p in blocking)
            + " — Schwellen vorsichtiger interpretieren."
        )
    return base


def render(
    result: AnalysisResult,
    plot_path: Path,
    out_path: Path,
    interpretation: dict[str, str] | None = None,
) -> Path:
    """Render the 5-page Leistungsdiagnostik report.

    Two-pass: (1) docxtpl fills text/KV-table/loop placeholders, (2) python-docx
    inserts the variable-width pivot table and colors the zone-table rows.
    Round 2 (Anna 2026-05-13) called for both features; docxtpl alone is awkward
    for variable column counts and per-row cell shading.
    """
    doc = DocxTemplate(_TEMPLATE_PATH)
    sport = result.test_run.athlete.sportart
    is_lauf = sport in {"lauf", "triathlon-lauf"}

    sport_labels = {
        "lauf": "Laufen",
        "rad": "Radfahren",
        "triathlon-rad": "Triathlon – Rad",
        "triathlon-lauf": "Triathlon – Lauf",
        "unspezifisch": "Unspezifisch",
    }
    intensity_units = {
        "lauf": "km/h",
        "rad": "W",
        "triathlon-rad": "W",
        "triathlon-lauf": "km/h",
        "unspezifisch": "Stufe",
    }
    v_max_labels = {
        "lauf": "Maximalgeschwindigkeit",
        "rad": "Maximalleistung",
        "triathlon-rad": "Maximalleistung",
        "triathlon-lauf": "Maximalgeschwindigkeit",
        "unspezifisch": "Maximalstufe",
    }
    intensity_unit = intensity_units[sport]
    v_max_label = v_max_labels[sport]
    x_axis_label = _x_axis_label(sport)

    def _round_intensity(v: float | None) -> float | str | None:
        if v is None:
            return None
        return round(v, 1) if is_lauf else round(v, 0)

    def _fmt_intensity(v: float | None) -> str:
        if v is None:
            return "—"
        r = _round_intensity(v)
        return f"{r}"

    def _render_zones() -> tuple:
        out = []
        for z in result.zones_final:
            if z.is_max_zone:
                # Round 3 (Anna 2026-05-17): Z6 → Intensität + Pace = "MAX",
                # HF-Feld leer. Frühere Variante zeigte "MAX" in allen drei
                # Feldern, das suggerierte aber eine bestimmbare Max-HF, die
                # in einem submaximalen Stufentest nicht vorliegt.
                v_range = pace_range = "MAX"
                hf_range = "—"
            elif z.is_collapsed_with_z2:
                # Round 3 (Anna 2026-05-17): Z1 == Z2 → Intensität, Pace, HF
                # leer lassen; nur RPE und Ziel bleiben aussagekräftig.
                v_range = hf_range = pace_range = "—"
            elif z.is_open_lower or (z.intensitaet_min is None and z.intensitaet_max is not None):
                hi = _round_intensity(z.intensitaet_max)
                v_range = f"≤ {hi} {intensity_unit}" if hi is not None else "—"
                hf_range = f"≤ {z.herzfrequenz_max}" if z.herzfrequenz_max else "—"
                pace_range = f"≥ {z.pace_min_min_per_km}" if z.pace_min_min_per_km else "—"
            else:
                lo = _round_intensity(z.intensitaet_min)
                hi = _round_intensity(z.intensitaet_max)
                v_range = f"{lo} – {hi}" if (lo is not None and hi is not None) else "—"
                if z.herzfrequenz_min and z.herzfrequenz_max:
                    hf_range = f"{z.herzfrequenz_min} – {z.herzfrequenz_max}"
                else:
                    hf_range = "—"
                if z.pace_max_min_per_km and z.pace_min_min_per_km:
                    pace_range = f"{z.pace_max_min_per_km} – {z.pace_min_min_per_km}"
                else:
                    pace_range = "—"
            rpe = (
                str(z.rpe_min) if z.rpe_min == z.rpe_max
                else f"{z.rpe_min} – {z.rpe_max}"
            )
            out.append({
                "name": z.name,
                "ziel": z.ziel,
                "intensitaet_range": v_range,
                "herzfrequenz_range": hf_range,
                "pace_range": pace_range,
                "rpe": rpe,
            })
        return tuple(out)

    def _render_steps() -> tuple:
        """Compact rows for the Page-2 Rohdaten table."""
        out = []
        for s in result.test_run.steps:
            out.append({
                "stufe": s.stufe,
                "intensitaet": _fmt_intensity(s.intensitaet),
                "herzfrequenz_bpm": s.herzfrequenz_bpm if s.herzfrequenz_bpm is not None else "—",
                "laktat_mmol": f"{s.laktat_mmol:.1f}" if s.laktat_mmol is not None else "—",
                "rpe": s.rpe if s.rpe is not None else "—",
            })
        return tuple(out)

    interp = interpretation or {}
    proto = result.test_run.testprotokoll
    athlete = result.test_run.athlete

    # Build the pivot data for the post-render step. Round 2 dropped rows with
    # no valid in-range root entirely. Round 3 (Anna 2026-05-17) adds a Word-
    # specific whitelist on top: 2.5 is hidden in the Word report (always),
    # 8.0 only appears when the cubic intersection lies inside the measured
    # range — i.e. the athlete actually reached that lactate level. JSON/Codex
    # keep ALL targets so trainers can still see + correct them.
    steps = result.test_run.steps
    measured_xs = [s.intensitaet for s in steps if s.laktat_mmol is not None]
    x_data_max = max(measured_xs) if measured_xs else float("inf")

    def _word_visible(row) -> bool:
        if row.intensitaet is None:
            return False
        if row.laktat not in _WORD_REPORT_LAKTAT_WHITELIST:
            return False
        if row.laktat in _WORD_REPORT_LAKTAT_REACHED_ONLY and row.intensitaet > x_data_max:
            return False
        return True

    valid_rows = [r for r in result.intersections if _word_visible(r)]
    intersection_pivot: dict[str, Any] = {
        "columns": [r.laktat for r in valid_rows],
        "rows": [
            {
                "label": x_axis_label,
                "values": [_fmt_intensity(r.intensitaet) for r in valid_rows],
            },
            {
                "label": "Pace (min/km)",
                "values": [(r.pace_min_per_km if r.pace_min_per_km else "—") for r in valid_rows],
            },
            {
                "label": "Herzfrequenz (bpm)",
                "values": [(str(r.herzfrequenz_bpm) if r.herzfrequenz_bpm is not None else "—") for r in valid_rows],
            },
        ],
    }

    context: dict[str, Any] = {
        "athlete": {
            "vorname": athlete.vorname,
            "name": athlete.name,
            "alter": athlete.alter if athlete.alter is not None else "—",
            "geburtsjahr": athlete.geburtsjahr if athlete.geburtsjahr is not None else "—",
            "email": athlete.email or "—",
            "gewicht_kg": athlete.gewicht_kg,
            "groesse_m": athlete.groesse_m,
            "geschlecht": athlete.geschlecht or "—",
            "sportart_label": sport_labels[sport],
            "trainingsziel": athlete.trainingsziel,
            "wettkampfziel": athlete.wettkampfziel,
            "trainingsumfang_wo": athlete.trainingsumfang_wo,
            "leistungsniveau": athlete.leistungsniveau,
        },
        "testprotokoll": {
            "testdatum": proto.testdatum.strftime("%d.%m.%Y"),
            "uhrzeit": proto.uhrzeit,
            "durchfuehrungsort": proto.durchfuehrungsort,
            "testleiter": proto.testleiter,
            "geraet": proto.geraet,
            "besonderheiten": proto.besonderheiten,
            "sportart_label": sport_labels[sport],
            "anfangsbelastung_display": f"{_round_intensity(proto.anfangsbelastung)} {intensity_unit}",
            "stufeninkrement_display": f"{_round_intensity(proto.stufeninkrement)} {intensity_unit}",
            "stufendauer_min": proto.stufendauer_min,
            "nachbelastungslaktat_3min": (
                f"{proto.nachbelastungslaktat_3min_mmol} mmol/l"
                if proto.nachbelastungslaktat_3min_mmol is not None else "—"
            ),
            "nachbelastungslaktat_5min": (
                f"{proto.nachbelastungslaktat_5min_mmol} mmol/l"
                if proto.nachbelastungslaktat_5min_mmol is not None else "—"
            ),
            # Round 3 (Anna 2026-05-17): Ruhelaktat, Steigung, Dauer letzte
            # Stufe als sichtbare Felder im Testprotokoll-KV. Leere Werte → `—`.
            "ruhelaktat_display": (
                f"{proto.ruhelaktat_mmol} mmol/l"
                if proto.ruhelaktat_mmol is not None else "—"
            ),
            "steigung_display": (
                f"{proto.steigung_prozent} %"
                if proto.steigung_prozent is not None else "—"
            ),
            "dauer_letzte_stufe_display": (
                f"{proto.dauer_letzte_stufe_min} min"
                if proto.dauer_letzte_stufe_min is not None
                else ("voll" if proto.letzte_stufe_vollstaendig else "—")
            ),
        },
        "v_max_label": v_max_label,
        "v_max_display": f"{_round_intensity(result.v_max)} {intensity_unit}",
        "ausbelastung_de": "ja" if proto.ausbelastung else "nein",
        "steps_display": _render_steps(),
        "zones": _render_zones(),
        # Round 3 P1-1 (Anna 2026-05-17): statische Trainingsformen-Mini-Tabelle
        # auf Seite 3 unter Trainingsbereiche. Ein Eintrag je Zone, dieselbe
        # Reihenfolge wie zones (Z1..Z6). Inhalt ist statisch, kommt aus
        # ZONE_METHODE in zones.py.
        "trainingsformen": tuple(
            {"name": z_name, "methode": ZONE_METHODE[z_name]}
            for z_name in ("Z1", "Z2", "Z3", "Z4", "Z5", "Z6")
        ),
        "diagram": InlineImage(doc, str(plot_path), width=Cm(_PLOT_WIDTH_CM)),
        "x_axis_label": x_axis_label,
        "internal_failed_checks": [p for p in result.pflichtpruefungen if not p.ok],
        "internal_passed_count": sum(1 for p in result.pflichtpruefungen if p.ok),
        "internal_total_count": len(result.pflichtpruefungen),
        # Pre-rendered text blocks for the Page-5 quadrants — docxtpl can't
        # safely nest `{%p if %}` + `{%p for %}` + `{%p else %}` + `{%p endif %}`
        # inside a single table cell without paragraph-boundary issues.
        "pflichtpruefungen_text": _format_pflichtpruefungen(result),
        "testqualitaet_text": _format_testqualitaet(result),
        # Coaching fields are truncated to per-field char caps so the
        # Page-5 Trainernotizen-Tabelle always fits in one page (Anna
        # 2026-05-17 — "immer 5 Seiten, immer gleiche Struktur"). Full
        # text remains in the JSON for archive.
        "coaching": {
            name: _truncate(getattr(result.test_run.coaching, name), cap, field=name)
            for name, cap in _COACHING_CHAR_CAPS.items()
        },
        # Page-4 interpretation blocks (Round 2 P1-5 — 4 sections).
        "interp_zusammenfassung": interp.get("zusammenfassung", "[Interpretation: ausstehend]"),
        "interp_schwellen": interp.get("schwellen", "[Interpretation: ausstehend]"),
        "interp_coaching_ausblick": interp.get(
            "coaching_ausblick_3_4_wochen",
            # Fallback to old "empfehlungen" key for backwards compat during migration.
            interp.get("empfehlungen", "[Interpretation: ausstehend]"),
        ),
        "interp_ernaehrung": interp.get("ernaehrung", "[Interpretation: ausstehend]"),
        # Page 5 internal narrative (optional).
        "interp_risiko": interp.get("risiko", None),
    }

    # Pass 1: docxtpl render into memory.
    doc.render(context)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    # Pass 2: python-docx post-processing — insert pivot table, color zone rows.
    final = Document(buf)
    _insert_intersection_pivot(final, intersection_pivot)
    _color_zone_rows(final)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    final.save(out_path)
    return out_path


# === Post-processing helpers ===

def _set_cell_bg(cell, hex_no_hash: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_no_hash)
    tc_pr.append(shd)


def _set_cell_borders(cell, color_hex: str = _LINE, size_pt: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size_pt))
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tc_pr.append(borders)


def _insert_intersection_pivot(doc, pivot: dict[str, Any]) -> None:
    """Find the `<<INTERSECTION_PIVOT>>` placeholder paragraph and replace it
    with a pivot table whose column count is the number of valid laktat values.
    """
    target = None
    for p in doc.paragraphs:
        if _PIVOT_MARKER in p.text:
            target = p
            break
    if target is None:
        return  # Template not updated yet — fail soft.

    cols = pivot["columns"]
    rows = pivot["rows"]

    if not cols:
        # No valid intersections at all — replace marker with explanatory text.
        target.text = (
            "Keine gültigen Schwellenschnittpunkte im Testbereich — Testqualität bitte prüfen."
        )
        return

    n_cols = len(cols) + 1  # +1 for the metric-label column
    n_rows = len(rows) + 1  # +1 for header

    new_table = doc.add_table(rows=n_rows, cols=n_cols)
    new_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_table.autofit = False

    # Header row: empty top-left cell, then Laktat-Wert headers.
    hdr_cells = new_table.rows[0].cells
    hdr_cells[0].text = ""
    _set_cell_borders(hdr_cells[0])
    hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = hdr_cells[0].paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run("Laktat (mmol/l)")
    run.font.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = _PRIMARY

    for j, lk in enumerate(cols, start=1):
        cell = hdr_cells[j]
        cell.text = ""
        _set_cell_bg(cell, "0B2545")
        _set_cell_borders(cell, "0B2545")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"{lk}")
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = _WHITE

    # Body rows: metric label on the left, values across.
    for i, row in enumerate(rows, start=1):
        cells = new_table.rows[i].cells
        # Label column.
        cells[0].text = ""
        _set_cell_borders(cells[0])
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(row["label"])
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = _PRIMARY
        # Value columns.
        for j, val in enumerate(row["values"], start=1):
            cell = cells[j]
            cell.text = ""
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = _TEXT

    # Move the table from its current position (appended at end of doc body)
    # to be right after the marker paragraph.
    target._element.addnext(new_table._element)
    # Remove the marker paragraph now that the table is in place.
    target._element.getparent().remove(target._element)


def _color_zone_rows(doc) -> None:
    """Find the Trainingsbereiche table and color each data row by zone."""
    # Identify the target table by header content: first row contains "Zone" + "Ziel".
    for table in doc.tables:
        if not table.rows:
            continue
        header_texts = [c.text.strip() for c in table.rows[0].cells]
        if "Zone" in header_texts and "Ziel" in header_texts:
            for row in table.rows[1:]:
                if not row.cells:
                    continue
                zone_name = row.cells[0].text.strip()
                bg = _ZONE_BG.get(zone_name)
                if not bg:
                    continue
                for cell in row.cells:
                    _set_cell_bg(cell, bg)
            return
